import fitz  # PyMuPDF
from docx import Document
import io
import re
from typing import Optional, Dict, List


class PDFProcessor:
    """Enhanced PDF and document processing using PyMuPDF."""
    
    def __init__(self):
        pass
    
    def extract_text_from_pdf(self, uploaded_file) -> Optional[str]:
        """
        Extract text from uploaded PDF file using PyMuPDF.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text as string or None if extraction fails
        """
        try:
            # Read the PDF file
            pdf_bytes = uploaded_file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            text = ""
            metadata = {}
            
            # Extract metadata
            metadata = doc.metadata
            
            # Extract text from all pages
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                # Clean page text
                page_text = self._clean_page_text(page_text)
                text += page_text + "\n"
            
            doc.close()
            
            # Final text cleaning
            text = self._clean_text(text)
            
            return text if text.strip() else None
            
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None
    
    def extract_text_with_formatting(self, uploaded_file) -> Dict:
        """
        Extract text with formatting information from PDF.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dictionary containing text, formatting, and metadata
        """
        try:
            pdf_bytes = uploaded_file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            result = {
                'text': '',
                'metadata': doc.metadata,
                'page_count': doc.page_count,
                'fonts': set(),
                'headings': [],
                'figures': [],
                'tables': []
            }
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                
                # Extract text with formatting
                blocks = page.get_text("dict")
                
                for block in blocks["blocks"]:
                    if "lines" in block:
                        for line in block["lines"]:
                            line_text = ""
                            for span in line["spans"]:
                                # Collect font information
                                font_info = f"{span['font']}_{span['size']}"
                                result['fonts'].add(font_info)
                                
                                # Identify potential headings (larger fonts)
                                if span['size'] > 14:
                                    result['headings'].append({
                                        'text': span['text'].strip(),
                                        'size': span['size'],
                                        'page': page_num + 1
                                    })
                                
                                line_text += span['text']
                            
                            if line_text.strip():
                                result['text'] += line_text + "\n"
                
                # Extract images/figures
                image_list = page.get_images()
                if image_list:
                    result['figures'].extend([{
                        'page': page_num + 1,
                        'count': len(image_list)
                    }])
            
            doc.close()
            
            # Clean the extracted text
            result['text'] = self._clean_text(result['text'])
            
            return result
            
        except Exception as e:
            print(f"Error extracting formatted text from PDF: {e}")
            return {'text': '', 'metadata': {}, 'page_count': 0, 'fonts': set(), 'headings': [], 'figures': [], 'tables': []}
    
    def extract_text_from_docx(self, uploaded_file) -> Optional[str]:
        """
        Extract text from uploaded DOCX file.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text as string or None if extraction fails
        """
        try:
            # Read the DOCX file
            doc = Document(io.BytesIO(uploaded_file.read()))
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            # Clean up the text
            text = self._clean_text(text)
            
            return text if text.strip() else None
            
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return None
    
    def _clean_page_text(self, text: str) -> str:
        """Clean text extracted from a single page."""
        if not text:
            return ""
        
        # Remove excessive whitespace within lines
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Clean line
            line = re.sub(r'\s+', ' ', line.strip())
            if line:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _clean_text(self, text: str) -> str:
        """
        Enhanced text cleaning and normalization.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers and headers/footers (enhanced patterns)
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        text = re.sub(r'\n\s*Page\s+\d+\s*\n', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'\n\s*\d+\s+of\s+\d+\s*\n', '\n', text, flags=re.IGNORECASE)
        
        # Clean up academic paper artifacts
        text = re.sub(r'\n\s*References?\s*\n', '\n\nReferences\n', text, flags=re.IGNORECASE)
        text = re.sub(r'\n\s*Bibliography\s*\n', '\n\nBibliography\n', text, flags=re.IGNORECASE)
        text = re.sub(r'\n\s*Abstract\s*\n', '\n\nAbstract\n', text, flags=re.IGNORECASE)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Clean up special characters
        text = text.replace('\x00', '')
        text = text.replace('\ufffd', '')
        text = text.replace('\f', '\n')  # Form feed to newline
        
        # Fix common PDF extraction issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between lowercase and uppercase
        text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)  # Add space after sentence endings
        
        return text.strip()
    
    def extract_sections(self, text: str) -> dict:
        """
        Enhanced section extraction with better patterns.
        
        Args:
            text: Full paper text
            
        Returns:
            Dictionary with section names as keys and content as values
        """
        sections = {}
        
        # Enhanced section patterns
        section_patterns = {
            'abstract': r'(?i)\babstract\b.*?(?=\n\s*(?:keywords?|introduction|1\.?\s*introduction|background|§|\Z))',
            'keywords': r'(?i)\bkeywords?\b.*?(?=\n\s*(?:introduction|1\.?\s*introduction|background|§|\Z))',
            'introduction': r'(?i)(?:introduction|1\.?\s*introduction).*?(?=\n\s*(?:related\s+work|literature\s+review|background|methodology|methods|2\.?\s*|§|\Z))',
            'literature_review': r'(?i)(?:related\s+work|literature\s+review|background|2\.?\s*(?:related|literature|background)).*?(?=\n\s*(?:methodology|methods|approach|3\.?\s*|§|\Z))',
            'methodology': r'(?i)(?:methodology|methods|approach|3\.?\s*(?:methodology|methods|approach)).*?(?=\n\s*(?:results|experiments?|evaluation|implementation|4\.?\s*|§|\Z))',
            'results': r'(?i)(?:results|experiments?|evaluation|4\.?\s*(?:results|experiments?|evaluation)).*?(?=\n\s*(?:discussion|analysis|conclusion|5\.?\s*|§|\Z))',
            'discussion': r'(?i)(?:discussion|analysis|5\.?\s*(?:discussion|analysis)).*?(?=\n\s*(?:conclusion|summary|6\.?\s*|§|\Z))',
            'conclusion': r'(?i)(?:conclusion|conclusions?|summary|6\.?\s*(?:conclusion|summary)).*?(?=\n\s*(?:references?|bibliography|acknowledgments?|appendix|§|\Z))',
            'references': r'(?i)(?:references?|bibliography).*?(?=\n\s*(?:appendix|§|\Z))',
            'acknowledgments': r'(?i)(?:acknowledgments?|acknowledgements?).*?(?=\n\s*(?:references?|bibliography|appendix|§|\Z))'
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.DOTALL)
            if match:
                content = match.group(0).strip()
                # Clean section title from content
                lines = content.split('\n')
                if len(lines) > 1:
                    content = '\n'.join(lines[1:]).strip()
                if content:
                    sections[section_name] = content
        
        return sections
    
    def count_citations(self, text: str) -> dict:
        """
        Enhanced citation analysis with multiple formats.
        
        Args:
            text: Paper text
            
        Returns:
            Dictionary with citation statistics
        """
        # Enhanced citation patterns
        citation_patterns = [
            r'\[(\d+(?:[-–,\s]*\d*)*)\]',  # [1], [2], [1-3], [1, 2, 3]
            r'\(([A-Z][a-zA-Z]+(?:\s+et\s+al\.?)?,?\s*\d{4}[a-z]?)\)',  # (Author 2020)
            r'\(([A-Z][a-zA-Z]+\s+(?:and|&)\s+[A-Z][a-zA-Z]+,?\s*\d{4})\)',  # (Author & Author, 2020)
            r'\(([A-Z][a-zA-Z]+(?:\s+et\s+al\.?)?,?\s*\d{4}[a-z]?(?:;\s*[A-Z][a-zA-Z]+(?:\s+et\s+al\.?)?,?\s*\d{4}[a-z]?)*)\)',  # Multiple citations
            r'([A-Z][a-zA-Z]+(?:\s+et\s+al\.?)?\s+\(\d{4}[a-z]?\))',  # Author (2020)
        ]
        
        all_citations = []
        citation_years = []
        
        for pattern in citation_patterns:
            matches = re.findall(pattern, text)
            all_citations.extend(matches)
            
            # Extract years from citations
            for match in matches:
                year_matches = re.findall(r'\b(19|20)\d{2}\b', str(match))
                citation_years.extend([int(year) for year in year_matches])
        
        # Remove duplicates
        unique_citations = list(set(all_citations))
        unique_years = list(set(citation_years))
        
        # Calculate recent citations (last 10 years)
        current_year = 2024
        recent_citations = [year for year in unique_years if year >= current_year - 10]
        
        # Calculate citation density per 1000 words
        word_count = len(text.split())
        citation_density = len(unique_citations) / max(word_count, 1) * 1000
        
        return {
            'total_citations': len(unique_citations),
            'unique_years': len(unique_years),
            'recent_citations': len(recent_citations),
            'citation_density': citation_density,
            'year_range': f"{min(unique_years)}-{max(unique_years)}" if unique_years else "N/A",
            'avg_year': sum(unique_years) / len(unique_years) if unique_years else 0
        }
    
    def extract_document_structure(self, text: str) -> Dict:
        """
        Analyze document structure and organization.
        
        Args:
            text: Full document text
            
        Returns:
            Dictionary with structure analysis
        """
        structure = {
            'total_words': len(text.split()),
            'total_sentences': len(re.findall(r'[.!?]+', text)),
            'total_paragraphs': len([p for p in text.split('\n\n') if p.strip()]),
            'sections': {},
            'figures_mentioned': 0,
            'tables_mentioned': 0,
            'equations_mentioned': 0
        }
        
        # Count figure, table, and equation references
        structure['figures_mentioned'] = len(re.findall(r'(?i)(?:figure|fig\.?)\s*\d+', text))
        structure['tables_mentioned'] = len(re.findall(r'(?i)table\s*\d+', text))
        structure['equations_mentioned'] = len(re.findall(r'(?i)(?:equation|eq\.?)\s*\d+', text))
        
        # Analyze sections
        sections = self.extract_sections(text)
        for section_name, content in sections.items():
            structure['sections'][section_name] = {
                'word_count': len(content.split()),
                'sentence_count': len(re.findall(r'[.!?]+', content)),
                'paragraph_count': len([p for p in content.split('\n\n') if p.strip()])
            }
        
        return structure